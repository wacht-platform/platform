from docx import Document
from docx.shared import Pt


def build_table_report(output_path: str) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)

    doc.add_heading("Vendor Comparison", level=0)
    doc.add_paragraph("This document compares candidate vendors across cost, fit, and delivery risk.")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Vendor"
    headers[1].text = "Annual Cost"
    headers[2].text = "Strength"
    headers[3].text = "Risk"

    for vendor, cost, strength, risk in [
        ("Vendor A", "$24k", "Fast setup", "Lower flexibility"),
        ("Vendor B", "$31k", "Best reporting", "Longer rollout"),
        ("Vendor C", "$19k", "Low cost", "Weaker controls"),
    ]:
        row = table.add_row().cells
        row[0].text = vendor
        row[1].text = cost
        row[2].text = strength
        row[3].text = risk

    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph("Choose the option that best balances delivery speed and control requirements.")
    doc.save(output_path)


if __name__ == "__main__":
    build_table_report("/tmp/vendor-comparison.docx")
