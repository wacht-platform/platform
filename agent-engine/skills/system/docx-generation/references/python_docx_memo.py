from docx import Document


def build_memo(output_path: str) -> None:
    doc = Document()
    doc.add_heading("Decision Memo", level=0)

    sections = [
        ("Summary", "Proceed with the phased rollout to reduce migration risk while preserving delivery momentum."),
        ("Context", "The team must choose between a single cutover and a staged migration."),
        ("Recommendation", "Adopt the staged path and review outcomes after the first milestone."),
        ("Implications", "This approach lowers immediate risk but extends the total transition timeline."),
    ]

    for heading, body in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)

    doc.save(output_path)


if __name__ == "__main__":
    build_memo("/tmp/decision-memo.docx")
