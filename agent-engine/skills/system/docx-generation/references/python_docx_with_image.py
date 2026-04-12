from docx import Document
from docx.shared import Inches


def build_illustrated_doc(output_path: str, image_path: str) -> None:
    doc = Document()
    doc.add_heading("System Overview", level=0)
    doc.add_paragraph("The image below summarizes the current architecture and data flow.")
    doc.add_picture(image_path, width=Inches(5.8))
    doc.add_heading("Notes", level=1)
    doc.add_paragraph("Review the boundary between ingestion and processing before the next rollout phase.")
    doc.save(output_path)


if __name__ == "__main__":
    build_illustrated_doc("/tmp/system-overview.docx", "/tmp/diagram.png")
