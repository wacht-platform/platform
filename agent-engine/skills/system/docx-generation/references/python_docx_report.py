from docx import Document
from docx.shared import Inches, Pt


def build_report(output_path: str) -> None:
    doc = Document()

    title = doc.add_heading("Market Analysis Report", level=0)
    title.alignment = 0

    intro = doc.add_paragraph()
    intro.add_run("Prepared for: ").bold = True
    intro.add_run("Internal Strategy Team")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report summarizes the market position, competitor landscape, and recommended next steps."
    )

    doc.add_heading("Key Findings", level=1)
    for item in [
        "Pricing remains the main competitive lever.",
        "The strongest differentiation is workflow depth, not breadth.",
        "Enterprise buyers care more about auditability than feature count.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Competitive Comparison", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Vendor"
    hdr[1].text = "Strength"
    hdr[2].text = "Risk"

    for row in [
        ("InboxDoctor", "Warmup depth", "Limited brand awareness"),
        ("Lemlist", "Distribution", "Generic positioning"),
        ("ZeroBounce", "Deliverability trust", "Less workflow focus"),
    ]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph(
        "Position the product around operational trust, measurable deliverability outcomes, and reusable execution playbooks."
    )

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)

    doc.save(output_path)


if __name__ == "__main__":
    build_report("/tmp/example_report.docx")
